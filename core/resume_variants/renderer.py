"""Safe renderers for generated resume variants."""

from __future__ import annotations

import html
import re
from io import BytesIO
from typing import Any

from docx import Document

MAX_RENDERED_BYTES = 256 * 1024


def safe_filename(stem: str, extension: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "-", stem).strip(".-") or "resume-variant"
    return f"{cleaned[:80]}.{extension}"


def _text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\x00", "").strip()


def _claims(content: dict[str, Any], key: str) -> list[dict[str, Any]]:
    value = content.get(key)
    return [claim for claim in value if isinstance(claim, dict)] if isinstance(value, list) else []


class ResumeVariantRenderer:
    """Render variant JSON without persisting generated binaries."""

    def render_markdown(self, content: dict[str, Any]) -> bytes:
        contact = _contact(content)
        lines = [f"# {_markdown_text(contact.get('name')) or 'Resume'}", ""]
        contact_line = _contact_line(contact, markdown=True)
        if contact_line:
            lines.extend([contact_line, ""])
        if _claims(content, "summary"):
            lines.append("## Professional Summary")
            for claim in _claims(content, "summary"):
                lines.append(_markdown_text(claim.get("text")))
            lines.append("")
        if _claims(content, "skills"):
            lines.append("## Skills")
            lines.append(", ".join(_markdown_text(claim.get("text")) for claim in _claims(content, "skills")))
            lines.append("")
        _append_markdown_entries(lines, content, "experience", "Experience")
        _append_markdown_entries(lines, content, "projects", "Projects")
        _append_markdown_education(lines, content)
        _append_markdown_simple_entries(lines, content, "certifications", "Certifications", _certification_text)
        _append_markdown_simple_entries(lines, content, "languages", "Languages", _language_text)
        return _bounded("\n".join(lines).encode("utf-8"))

    def render_html(self, content: dict[str, Any]) -> bytes:
        contact = _contact(content)
        sections = [
            "<!doctype html><html><head><meta charset=\"utf-8\">",
            "<title>Resume</title>",
            "<style>body{font-family:Arial,sans-serif;max-width:850px;margin:32px auto;line-height:1.35;color:#111}"
            "h1{margin-bottom:4px}h2{border-bottom:1px solid #777;padding-bottom:3px;margin-top:22px}"
            "h3{margin-bottom:2px}.meta{color:#444}.dates{float:right;font-weight:normal}li{margin:3px 0}</style>",
            "</head><body>",
            f"<h1>{html.escape(_text(contact.get('name')) or 'Resume')}</h1>",
        ]
        contact_line = _contact_line(contact)
        if contact_line:
            sections.append(f"<div class=\"meta\">{html.escape(contact_line)}</div>")
        if _claims(content, "summary"):
            sections.append("<h2>Professional Summary</h2>")
            for claim in _claims(content, "summary"):
                sections.append(f"<p>{html.escape(_text(claim.get('text')))}</p>")
        if _claims(content, "skills"):
            skills = ", ".join(html.escape(_text(claim.get("text"))) for claim in _claims(content, "skills"))
            sections.append(f"<h2>Skills</h2><p>{skills}</p>")
        _append_html_entries(sections, content, "experience", "Experience")
        _append_html_entries(sections, content, "projects", "Projects")
        _append_html_education(sections, content)
        _append_html_simple_entries(sections, content, "certifications", "Certifications", _certification_text)
        _append_html_simple_entries(sections, content, "languages", "Languages", _language_text)
        sections.append("</body></html>")
        return _bounded("".join(sections).encode("utf-8"))

    def render_docx(self, content: dict[str, Any]) -> bytes:
        document = Document()
        contact = _contact(content)
        document.add_heading(_text(contact.get("name")) or "Resume", level=1)
        contact_line = _contact_line(contact)
        if contact_line:
            document.add_paragraph(contact_line)
        if _claims(content, "summary"):
            document.add_heading("Professional Summary", level=2)
            for claim in _claims(content, "summary"):
                document.add_paragraph(_text(claim.get("text")))
        if _claims(content, "skills"):
            document.add_heading("Skills", level=2)
            document.add_paragraph(", ".join(_text(claim.get("text")) for claim in _claims(content, "skills")))
        _append_docx_entries(document, content, "experience", "Experience")
        _append_docx_entries(document, content, "projects", "Projects")
        _append_docx_education(document, content)
        _append_docx_simple_entries(document, content, "certifications", "Certifications", _certification_text)
        _append_docx_simple_entries(document, content, "languages", "Languages", _language_text)
        output = BytesIO()
        document.save(output)
        return _bounded(output.getvalue())

def _contact(content: dict[str, Any]) -> dict[str, Any]:
    value = content.get("contact")
    return value if isinstance(value, dict) else {}

def _contact_line(contact: dict[str, Any], *, markdown: bool = False) -> str:
    values = [_text(contact.get(key)) for key in ("email", "phone", "location")]
    links = contact.get("links")
    if isinstance(links, list):
        values.extend(_text(link) for link in links[:8])
    cleaned = [value for value in values if value]
    if markdown:
        cleaned = [_markdown_text(value) for value in cleaned]
    return " | ".join(cleaned)

def _entry_heading(entry: dict[str, Any], section: str) -> str:
    keys = ("title", "company") if section == "experience" else ("name",)
    return " - ".join(_text(entry.get(key)) for key in keys if _text(entry.get(key)))

def _entry_dates(entry: dict[str, Any], section: str) -> str:
    if section == "experience":
        return " – ".join(
            value for value in (_text(entry.get("start_date")), _text(entry.get("end_date"))) if value
        )
    return _text(entry.get("date"))

def _append_markdown_entries(lines: list[str], content: dict[str, Any], key: str, heading: str) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    lines.append(f"## {heading}")
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        title = _markdown_text(_entry_heading(entry, key))
        dates = _markdown_text(_entry_dates(entry, key))
        if title:
            lines.append(f"### {title}{f' | {dates}' if dates else ''}")
        technologies = entry.get("technologies")
        if isinstance(technologies, list) and technologies:
            lines.append(f"*{', '.join(_markdown_text(value) for value in technologies if _text(value))}*")
        for bullet in _claims(entry, "bullets"):
            lines.append(f"- {_markdown_text(bullet.get('text'))}")
        lines.append("")

def _append_markdown_education(lines: list[str], content: dict[str, Any]) -> None:
    entries = content.get("education")
    if not isinstance(entries, list) or not entries:
        return
    lines.append("## Education")
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        degree = " in ".join(part for part in (_text(entry.get("degree")), _text(entry.get("field_of_study"))) if part)
        heading = " - ".join(part for part in (degree, _text(entry.get("institution"))) if part)
        year = _text(entry.get("graduation_year"))
        if heading:
            lines.append(f"### {_markdown_text(heading)}{f' | {_markdown_text(year)}' if year else ''}")
        for detail in _claims(entry, "details"):
            lines.append(f"- {_markdown_text(detail.get('text'))}")
        lines.append("")

def _append_markdown_simple_entries(lines, content, key, heading, formatter) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    lines.append(f"## {heading}")
    lines.extend(f"- {_markdown_text(formatter(entry))}" for entry in entries if isinstance(entry, dict) and formatter(entry))
    lines.append("")

def _append_html_entries(sections: list[str], content: dict[str, Any], key: str, heading: str) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    sections.append(f"<h2>{heading}</h2>")
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        title = _entry_heading(entry, key)
        dates = _entry_dates(entry, key)
        if title:
            date_html = f'<span class="dates">{html.escape(dates)}</span>' if dates else ""
            sections.append(f"<h3>{html.escape(title)}{date_html}</h3>")
        technologies = entry.get("technologies")
        if isinstance(technologies, list) and technologies:
            sections.append(f"<div class=\"meta\">{html.escape(', '.join(_text(value) for value in technologies if _text(value)))}</div>")
        bullets = _claims(entry, "bullets")
        if bullets:
            sections.append("<ul>")
            sections.extend(f"<li>{html.escape(_text(bullet.get('text')))}</li>" for bullet in bullets)
            sections.append("</ul>")

def _append_html_education(sections: list[str], content: dict[str, Any]) -> None:
    entries = content.get("education")
    if not isinstance(entries, list) or not entries:
        return
    sections.append("<h2>Education</h2>")
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        degree = " in ".join(part for part in (_text(entry.get("degree")), _text(entry.get("field_of_study"))) if part)
        heading = " - ".join(part for part in (degree, _text(entry.get("institution"))) if part)
        year = _text(entry.get("graduation_year"))
        if heading:
            date_html = f'<span class="dates">{html.escape(year)}</span>' if year else ""
            sections.append(f"<h3>{html.escape(heading)}{date_html}</h3>")
        details = _claims(entry, "details")
        if details:
            sections.append("<ul>")
            sections.extend(f"<li>{html.escape(_text(detail.get('text')))}</li>" for detail in details)
            sections.append("</ul>")

def _append_html_simple_entries(sections, content, key, heading, formatter) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    sections.append(f"<h2>{heading}</h2><ul>")
    sections.extend(f"<li>{html.escape(formatter(entry))}</li>" for entry in entries if isinstance(entry, dict) and formatter(entry))
    sections.append("</ul>")

def _append_docx_entries(document, content: dict[str, Any], key: str, heading: str) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    document.add_heading(heading, level=2)
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        title = _entry_heading(entry, key)
        dates = _entry_dates(entry, key)
        if title:
            document.add_heading(f"{title}{f' | {dates}' if dates else ''}", level=3)
        technologies = entry.get("technologies")
        if isinstance(technologies, list) and technologies:
            document.add_paragraph(", ".join(_text(value) for value in technologies if _text(value)))
        for bullet in _claims(entry, "bullets"):
            document.add_paragraph(_text(bullet.get("text")), style="List Bullet")

def _append_docx_education(document, content: dict[str, Any]) -> None:
    entries = content.get("education")
    if not isinstance(entries, list) or not entries:
        return
    document.add_heading("Education", level=2)
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        degree = " in ".join(part for part in (_text(entry.get("degree")), _text(entry.get("field_of_study"))) if part)
        heading = " - ".join(part for part in (degree, _text(entry.get("institution"))) if part)
        year = _text(entry.get("graduation_year"))
        if heading:
            document.add_heading(f"{heading}{f' | {year}' if year else ''}", level=3)
        for detail in _claims(entry, "details"):
            document.add_paragraph(_text(detail.get("text")), style="List Bullet")

def _append_docx_simple_entries(document, content, key, heading, formatter) -> None:
    entries = content.get(key)
    if not isinstance(entries, list) or not entries:
        return
    document.add_heading(heading, level=2)
    for entry in entries:
        if isinstance(entry, dict) and formatter(entry):
            document.add_paragraph(formatter(entry), style="List Bullet")

def _certification_text(entry: dict[str, Any]) -> str:
    name = _text(entry.get("name"))
    issuer = _text(entry.get("issuer"))
    issued = _text(entry.get("issued_year"))
    return " - ".join(part for part in (name, issuer, issued) if part)

def _language_text(entry: dict[str, Any]) -> str:
    language = _text(entry.get("language"))
    proficiency = _text(entry.get("proficiency"))
    return " - ".join(part for part in (language, proficiency) if part)


def _markdown_text(value: Any) -> str:
    return html.escape(_text(value), quote=False)


def _bounded(payload: bytes) -> bytes:
    if len(payload) > MAX_RENDERED_BYTES:
        raise ValueError("Rendered resume variant exceeds size limit.")
    return payload
