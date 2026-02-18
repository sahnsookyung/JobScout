"""Tests for the multi-format resume parser."""
import json
import tempfile
from pathlib import Path
from unittest import TestCase

import yaml
from docx import Document
from pypdf import PdfWriter

from etl.resume.parser import ResumeParser, ParsedResume


class TestResumeParser(TestCase):
    """Test suite for ResumeParser."""

    def setUp(self):
        """Set up test fixtures."""
        self.parser = ResumeParser()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _create_temp_file(self, filename: str, content: str | bytes) -> str:
        """Helper to create a temp file."""
        path = Path(self.temp_dir) / filename
        mode = 'wb' if isinstance(content, bytes) else 'w'
        with open(path, mode) as f:
            f.write(content)
        return str(path)

    def test_parse_json_valid(self):
        """Should parse valid JSON resume."""
        data = {
            "name": "John Doe",
            "title": "Software Engineer",
            "sections": [
                {
                    "title": "Experience",
                    "items": [
                        {"company": "Tech Corp", "role": "Developer", "period": "2020-2023"}
                    ]
                }
            ]
        }
        path = self._create_temp_file('resume.json', json.dumps(data))

        result = self.parser.parse(path)

        self.assertIsNotNone(result.data)
        self.assertEqual(result.format, 'json')
        self.assertEqual(result.data['name'], 'John Doe')
        self.assertIn('John Doe', result.text)

    def test_parse_json_invalid(self):
        """Should raise ValueError for invalid JSON."""
        path = self._create_temp_file('invalid.json', '{invalid json}')

        with self.assertRaises(ValueError) as ctx:
            self.parser.parse(path)

        self.assertIn('Invalid JSON', str(ctx.exception))

    def test_parse_yaml_valid(self):
        """Should parse valid YAML resume."""
        content = """
name: Jane Smith
title: Data Scientist
experience:
  - company: Data Corp
    role: Analyst
    years: 2
"""
        path = self._create_temp_file('resume.yaml', content)

        result = self.parser.parse(path)

        self.assertIsNotNone(result.data)
        self.assertEqual(result.format, 'yaml')
        self.assertEqual(result.data['name'], 'Jane Smith')

    def test_parse_txt(self):
        """Should parse plain text resume."""
        content = """John Doe
Software Engineer

Experience:
- 5 years Python development
- Backend systems design"""
        path = self._create_temp_file('resume.txt', content)

        result = self.parser.parse(path)

        self.assertIsNone(result.data)
        self.assertEqual(result.format, 'txt')
        self.assertIn('Software Engineer', result.text)

    def test_parse_docx(self):
        """Should parse Word document resume."""
        path = Path(self.temp_dir) / 'resume.docx'
        doc = Document()
        doc.add_heading('Jane Doe', 0)
        doc.add_paragraph('Senior Developer')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('5 years at Tech Corp')
        doc.save(path)

        result = self.parser.parse(str(path))

        self.assertIsNone(result.data)
        self.assertEqual(result.format, 'docx')
        self.assertIn('Jane Doe', result.text)
        self.assertIn('Tech Corp', result.text)

    def test_parse_pdf(self):
        """Should parse PDF resume."""
        # Create a simple PDF with text
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        import io

        # Create PDF in memory
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        c.drawString(100, 700, "John Developer")
        c.drawString(100, 680, "Full Stack Engineer")
        c.drawString(100, 650, "Experience:")
        c.drawString(120, 630, "- Python Developer at TechCorp")
        c.save()
        packet.seek(0)

        path = Path(self.temp_dir) / 'resume.pdf'
        with open(path, 'wb') as f:
            f.write(packet.getvalue())

        result = self.parser.parse(str(path))

        self.assertIsNone(result.data)
        self.assertEqual(result.format, 'pdf')
        self.assertIn('John Developer', result.text)

    def test_parse_file_not_found(self):
        """Should raise FileNotFoundError for missing file."""
        with self.assertRaises(FileNotFoundError):
            self.parser.parse('/nonexistent/resume.pdf')

    def test_parse_unsupported_format(self):
        """Should raise ValueError for unsupported format."""
        path = self._create_temp_file('resume.rtf', 'Some content')

        with self.assertRaises(ValueError) as ctx:
            self.parser.parse(path)

        self.assertIn('Unsupported', str(ctx.exception))

    def test_is_supported(self):
        """Should correctly check if format is supported."""
        self.assertTrue(self.parser.is_supported('resume.json'))
        self.assertTrue(self.parser.is_supported('resume.pdf'))
        self.assertTrue(self.parser.is_supported('resume.docx'))
        self.assertFalse(self.parser.is_supported('resume.exe'))
        self.assertFalse(self.parser.is_supported('resume'))

    def test_get_supported_formats(self):
        """Should return list of supported formats."""
        formats = ResumeParser.get_supported_formats()

        self.assertIn('.json', formats)
        self.assertIn('.pdf', formats)
        self.assertIn('.docx', formats)
        self.assertIn('.yaml', formats)
        self.assertIn('.yml', formats)
        self.assertIn('.txt', formats)

    def test_parse_json_with_unicode(self):
        """Should handle Unicode characters in JSON."""
        data = {"name": "山田太郎", "title": "エンジニア"}
        path = self._create_temp_file('unicode.json', json.dumps(data, ensure_ascii=False))

        result = self.parser.parse(path)

        self.assertEqual(result.data['name'], '山田太郎')
        self.assertIn('山田太郎', result.text)

    def test_parse_empty_txt(self):
        """Should raise ValueError for empty text file."""
        path = self._create_temp_file('empty.txt', '')

        with self.assertRaises(ValueError) as ctx:
            self.parser.parse(path)

        self.assertIn('Empty resume file', str(ctx.exception))

    def test_parse_yaml_not_dict(self):
        """Should raise ValueError if YAML doesn't contain a dict."""
        content = "- item1\n- item2"  # A list, not a dict
        path = self._create_temp_file('list.yaml', content)

        with self.assertRaises(ValueError) as ctx:
            self.parser.parse(path)

        self.assertIn('mapping/dict', str(ctx.exception))
