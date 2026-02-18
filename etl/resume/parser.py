"""
Multi-format Resume Parser - Parse resumes from various formats.

Supports:
- JSON (.json): Returns dict + JSON-as-text for LLM
- YAML (.yaml, .yml): Returns dict + text for LLM
- Plain Text (.txt): Returns None dict + text for LLM
- Word Documents (.docx): Returns None dict + extracted text for LLM
- PDF (.pdf): Returns None dict + extracted text for LLM

For JSON/YAML formats, returns structured dict that can be used directly.
For other formats, returns raw text that needs LLM extraction.
"""
import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Dict, Any

import yaml
from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)


@dataclass
class ParsedResume:
    """Result of parsing a resume file.

    Attributes:
        data: Structured data dict for JSON/YAML formats, None for text-based formats
        text: Extracted text suitable for LLM processing
        format: Detected file format (e.g., 'json', 'yaml', 'txt', 'docx', 'pdf')
        source_path: Original file path
    """
    data: Optional[Dict[str, Any]]
    text: str
    format: str
    source_path: str


class ResumeParser:
    """Parse resumes from multiple formats.

    Handles format detection and extraction. For structured formats (JSON/YAML),
    returns both parsed data and text representation. For document formats
    (PDF, DOCX, TXT), extracts text for LLM-based parsing.
    """

    SUPPORTED_FORMATS = {
        '.json', '.yaml', '.yml', '.txt', '.docx', '.pdf'
    }

    def __init__(self):
        """Initialize the resume parser."""
        self.logger = logging.getLogger(__name__)

    def parse(self, file_path: str) -> ParsedResume:
        """Parse a resume file and extract content.

        Detects format from file extension and routes to appropriate
        parser. Returns structured data for JSON/YAML, raw text for
        document formats.

        Args:
            file_path: Path to the resume file

        Returns:
            ParsedResume with extracted data/text

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If format is unsupported or parsing fails
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_FORMATS:
            supported = ', '.join(sorted(self.SUPPORTED_FORMATS))
            raise ValueError(
                f"Unsupported resume format: {ext}. "
                f"Supported formats: {supported}"
            )

        self.logger.info(f"Parsing resume from {file_path} (format: {ext})")

        # Route to appropriate parser
        if ext == '.json':
            return self._parse_json(path)
        elif ext in ('.yaml', '.yml'):
            return self._parse_yaml(path)
        elif ext == '.txt':
            return self._parse_txt(path)
        elif ext == '.docx':
            return self._parse_docx(path)
        elif ext == '.pdf':
            return self._parse_pdf(path)

        # Should never reach here due to extension check above
        raise ValueError(f"Unexpected format: {ext}")

    def _parse_json(self, path: Path) -> ParsedResume:
        """Parse JSON resume file.

        Returns both the parsed dict and JSON string representation.
        The dict can be used directly, while the text is for LLM context.

        Args:
            path: Path to JSON file

        Returns:
            ParsedResume with dict data and JSON text
        """
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Convert back to formatted JSON for LLM text
            text = json.dumps(data, indent=2, ensure_ascii=False)

            self.logger.debug(f"Parsed JSON resume from {path}")

            return ParsedResume(
                data=data,
                text=text,
                format='json',
                source_path=str(path)
            )

        except json.JSONDecodeError as e:
            raise ValueError(
                f"Invalid JSON in resume file: {e}. "
                f"Check line {e.lineno}, column {e.colno}"
            )
        except UnicodeDecodeError as e:
            raise ValueError(
                f"File encoding issue in {path}: {e}. "
                f"Ensure file is UTF-8 encoded."
            )

    def _parse_yaml(self, path: Path) -> ParsedResume:
        """Parse YAML resume file.

        Returns both the parsed dict and formatted text representation.

        Args:
            path: Path to YAML file

        Returns:
            ParsedResume with dict data and YAML text
        """
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)

            # Validate it's a dict (YAML can load any type)
            if not isinstance(data, dict):
                raise ValueError(
                    f"YAML resume must contain a mapping/dict, got {type(data).__name__}"
                )

            # Convert to JSON for consistent LLM text format
            text = json.dumps(data, indent=2, ensure_ascii=False)

            self.logger.debug(f"Parsed YAML resume from {path}")

            return ParsedResume(
                data=data,
                text=text,
                format='yaml',
                source_path=str(path)
            )

        except yaml.YAMLError as e:
            raise ValueError(f"Invalid YAML in resume file: {e}")
        except UnicodeDecodeError as e:
            raise ValueError(
                f"File encoding issue in {path}: {e}. "
                f"Ensure file is UTF-8 encoded."
            )

    def _parse_txt(self, path: Path) -> ParsedResume:
        """Parse plain text resume file.

        Returns raw text content. No structured data available.

        Args:
            path: Path to text file

        Returns:
            ParsedResume with None data and text content
        """
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()

            if not text.strip():
                raise ValueError(f"Empty resume file: {path}")

            return ParsedResume(
                data=None,
                text=text,
                format='txt',
                source_path=str(path)
            )

        except UnicodeDecodeError as e:
            raise ValueError(
                f"File encoding issue in {path}: {e}. "
                f"Ensure file is UTF-8 encoded."
            )

    def _parse_docx(self, path: Path) -> ParsedResume:
        """Parse Word document (.docx) resume.

        Extracts text from all paragraphs. No structured data available.

        Args:
            path: Path to DOCX file

        Returns:
            ParsedResume with None data and extracted text
        """
        try:
            doc = Document(str(path))

            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Also extract from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraphs.append(' '.join(row_texts))

            text = '\n\n'.join(paragraphs)

            if not text.strip():
                self.logger.warning(f"Empty or minimal content in DOCX: {path}")

            self.logger.debug(f"Parsed DOCX resume from {path}")

            return ParsedResume(
                data=None,
                text=text,
                format='docx',
                source_path=str(path)
            )

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file {path}: {e}")

    def _parse_pdf(self, path: Path) -> ParsedResume:
        """Parse PDF resume file.

        Extracts text from all pages. No structured data available.

        Args:
            path: Path to PDF file

        Returns:
            ParsedResume with None data and extracted text
        """
        try:
            reader = PdfReader(path)

            if len(reader.pages) == 0:
                raise ValueError("PDF file has no pages")

            # Extract text from all pages
            pages_text = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages_text.append(page_text.strip())
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {i + 1}: {e}")

            text = '\n\n'.join(pages_text)

            if not text.strip():
                self.logger.warning(
                    f"No text extracted from PDF {path}. "
                    f"The PDF may be scanned images or have text extraction disabled."
                )

            self.logger.debug(
                f"Parsed PDF resume from {path} ({len(reader.pages)} pages, "
                f"{len(text)} chars extracted)"
            )

            return ParsedResume(
                data=None,
                text=text,
                format='pdf',
                source_path=str(path)
            )

        except Exception as e:
            raise ValueError(f"Failed to parse PDF file {path}: {e}")

    def is_supported(self, file_path: str) -> bool:
        """Check if a file format is supported.

        Args:
            file_path: Path to check

        Returns:
            True if format is supported, False otherwise
        """
        return Path(file_path).suffix.lower() in self.SUPPORTED_FORMATS

    @classmethod
    def get_supported_formats(cls) -> list[str]:
        """Get list of supported file extensions.

        Returns:
            List of supported file extensions (e.g., ['.json', '.pdf'])
        """
        return sorted(cls.SUPPORTED_FORMATS)
